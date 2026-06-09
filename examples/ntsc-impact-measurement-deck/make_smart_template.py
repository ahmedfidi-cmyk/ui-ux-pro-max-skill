#!/usr/bin/env python3
"""NTSC SMART Goal Template for event participation — fillable, RTL Arabic, branded."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EVENING="004E43"; GABLE="17332F"; HIGHLAND="71946A"; ANZAC="DEB83B"
SWIRL="D5D0C3"; PAPER="F6F4EE"; INK="2C3A36"; WHITE="FFFFFF"; LIGHT="EEF2EC"
AR="Segoe UI"   # swap to brand "TheSans" where installed

doc=Document()
# ---- page setup A4 + margins ----
sec=doc.sections[0]
sec.page_height=Mm(297); sec.page_width=Mm(210)
sec.top_margin=Cm(1.4); sec.bottom_margin=Cm(1.4)
sec.left_margin=Cm(1.6); sec.right_margin=Cm(1.6)

def _rtl_par(p):
    pPr=p._p.get_or_add_pPr()
    b=OxmlElement('w:bidi'); pPr.append(b)
def _shade(el,color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color)
    el.append(sh)
def _set_cell_bg(cell,color):
    _shade(cell._tc.get_or_add_tcPr(),color)
def _cell_margins(cell,t=60,b=60,l=100,r=100):
    tcPr=cell._tc.get_or_add_tcPr()
    m=OxmlElement('w:tcMar')
    for tag,v in (('top',t),('bottom',b),('start',l),('end',r)):
        e=OxmlElement('w:'+tag); e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)
def _valign(cell,v='center'):
    tcPr=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:vAlign'); e.set(qn('w:val'),v); tcPr.append(e)

def run(p,text,size=11,color=INK,bold=False,italic=False,font=AR):
    r=p.add_run(text); f=r.font
    f.size=Pt(size); f.bold=bold; f.italic=italic; f.name=font
    f.color.rgb=RGBColor.from_string(color)
    rPr=r._r.get_or_add_rPr()
    rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:cs'),font); rf.set(qn('w:ascii'),font); rf.set(qn('w:hAnsi'),font)
    rtl=OxmlElement('w:rtl'); rtl.set(qn('w:val'),'1'); rPr.append(rtl)
    return r

def para(text="",size=11,color=INK,bold=False,align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_before=0,space_after=4,italic=False,font=AR,line=1.2):
    p=doc.add_paragraph(); p.alignment=align; _rtl_par(p)
    pf=p.paragraph_format
    pf.space_before=Pt(space_before); pf.space_after=Pt(space_after); pf.line_spacing=line
    if text: run(p,text,size,color,bold,italic,font)
    return p

def hrule(color=ANZAC,size=14):
    p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),str(size))
    bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),color)
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(0)

def section_head(ar,en,num):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
    set_table_width(t,[18.8]); _table_rtl(t)
    c=t.rows[0].cells[0]; _set_cell_bg(c,EVENING); _cell_margins(c,90,90,140,140)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
    run(p,f"{num}  ",12,ANZAC,True); run(p,ar,13,WHITE,True)
    run(p,f"   ·  {en}",9,"CFE0D9",False)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def _table_rtl(t):
    tblPr=t._tbl.tblPr
    bv=OxmlElement('w:bidiVisual'); tblPr.append(bv)
def _set_borders(t,color="CFDAD2",sz=6):
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for tag in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+tag); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color); b.append(e)
    tblPr.append(b)
def set_table_width(t,widths_cm):
    t.autofit=False
    tblPr=t._tbl.tblPr
    layout=OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); tblPr.append(layout)
    grid=t._tbl.find(qn('w:tblGrid'))
    for i,w in enumerate(widths_cm):
        for cell in t.columns[i].cells:
            cell.width=Cm(w)

# ================= HEADER (logo) =================
htab=doc.add_table(rows=1,cols=2); htab.alignment=WD_TABLE_ALIGNMENT.CENTER
set_table_width(htab,[12.8,6.0]); _table_rtl(htab)
# right cell = title block, left = logo  (RTL so first col renders right)
ctitle=htab.rows[0].cells[0]; clogo=htab.rows[0].cells[1]
_valign(ctitle); _valign(clogo)
p=ctitle.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
run(p,"نموذج أهداف ",17,EVENING,True); run(p,"SMART",17,ANZAC,True)
p2=ctitle.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p2)
run(p2,"لمشاركات المركز في الفعاليات والمؤتمرات",13,EVENING,True)
p3=ctitle.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p3)
run(p3,"SMART Goals Template · Event Participation",9,HIGHLAND,False)
# logo
lp=clogo.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
try:
    lp.add_run().add_picture("assets/ntsc-logo-green.png", width=Cm(5.0))
except Exception as e:
    run(lp,"NTSC",18,EVENING,True)
doc.add_paragraph().paragraph_format.space_after=Pt(2)
hrule(ANZAC,18)
para("إدارة التواصل · المركز الوطني لسلامة النقل   |   Communications Department — National Transport Safety Center",
     8.5,HIGHLAND,False,WD_ALIGN_PARAGRAPH.RIGHT,space_after=8)

# intro
para("يُستخدم هذا النموذج في مرحلة التخطيط (قبل الحدث) لصياغة أهداف المشاركة وفق منهجية SMART، "
     "وربطها بمؤشرات الأداء المعتمدة لقياس الأثر بعد المشاركة.",11,INK,False,space_after=8)

# ================= 1. EVENT DATA =================
section_head("بيانات الفعالية","Event Data","1")
fields=[("اسم الفعالية","Event name"),("الجهة المنظِّمة","Organizer"),
        ("مكان وتاريخ المشاركة","Venue & date"),("نوع المشاركة","Participation type"),
        ("مدير المشاركة","Lead"),("فريق ممثّلي المركز","Representatives")]
t=doc.add_table(rows=len(fields),cols=2); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
set_table_width(t,[13.8,5.0]); _table_rtl(t); _set_borders(t)
for i,(ar,en) in enumerate(fields):
    lbl=t.rows[i].cells[1]; val=t.rows[i].cells[0]
    _set_cell_bg(lbl,LIGHT); _cell_margins(lbl); _cell_margins(val); _valign(lbl); _valign(val)
    p=lbl.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
    run(p,ar,10.5,EVENING,True); run(p,f"  ·  {en}",8,HIGHLAND)
    vp=val.paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(vp)
    run(vp," ",10.5)  # blank to fill
doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ================= 2. OVERALL OBJECTIVE =================
section_head("الهدف العام من المشاركة","Overall Objective","2")
t=doc.add_table(rows=1,cols=1); set_table_width(t,[18.8]); _table_rtl(t); _set_borders(t)
c=t.rows[0].cells[0]; _cell_margins(c,120,260,140,140)
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
run(p,"اكتب هنا الغاية الإستراتيجية من المشاركة… ",10.5,"9AA7A2",italic=True)
doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ================= 3. SMART FRAMEWORK =================
section_head("إطار SMART لصياغة الهدف","The SMART Framework","3")
smart=[("S","محدد","Specific","ماذا نريد تحقيقه بالضبط؟ مَن المعنيّ وأين؟","E1B92E"),
       ("M","قابل للقياس","Measurable","ما المؤشر والرقم المستهدف الذي يثبت تحقّق الهدف؟",EVENING),
       ("A","قابل للتحقيق","Achievable","هل الموارد (فريق، ميزانية، وقت) كافية لتحقيقه؟",HIGHLAND),
       ("R","ذو صلة","Relevant","كيف يخدم الهدفُ أولويات المركز وصورته الذهنية؟",GABLE),
       ("T","محدد زمنياً","Time-bound","ما الإطار الزمني والموعد النهائي للقياس؟",ANZAC)]
t=doc.add_table(rows=len(smart)+1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
set_table_width(t,[10.8,6.2,1.8]); _table_rtl(t); _set_borders(t)
hdr=["السؤال التوجيهي · Guiding question","المعنى · Meaning","الحرف"]
for j,htext in enumerate(hdr):
    cc=t.rows[0].cells[j]; _set_cell_bg(cc,EVENING); _cell_margins(cc,70,70,120,120); _valign(cc)
    pp=cc.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl_par(pp)
    run(pp,htext,9.5,WHITE,True)
for i,(L,ar,en,q,col) in enumerate(smart,1):
    letter=t.rows[i].cells[2]; mean=t.rows[i].cells[1]; ques=t.rows[i].cells[0]
    _set_cell_bg(letter,col); _valign(letter); _cell_margins(letter,90,90,60,60)
    lp=letter.paragraphs[0]; lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(lp,L,16,WHITE if col!=ANZAC else GABLE,True,font="Segoe UI")
    _cell_margins(mean); _valign(mean)
    mp=mean.paragraphs[0]; mp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(mp)
    run(mp,ar,11,EVENING,True); run(mp,f"  {en}",8.5,HIGHLAND)
    _cell_margins(ques); _valign(ques)
    qp=ques.paragraphs[0]; qp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(qp)
    run(qp,q,10,INK)
doc.add_paragraph().paragraph_format.space_after=Pt(6)

# ================= 4. SMART GOAL WORKSHEET (fillable, repeatable) =================
section_head("بطاقة صياغة الهدف (قابلة للتكرار لكل هدف)","SMART Goal Worksheet","4")
def goal_card(title_default=""):
    t=doc.add_table(rows=6,cols=2); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
    set_table_width(t,[15.8,3.0]); _table_rtl(t); _set_borders(t)
    rows=[("الهدف (صياغة نهائية)","Goal statement",title_default),
          ("S — محدد","Specific",""),
          ("M — قابل للقياس (المؤشر + الرقم المستهدف)","Measurable",""),
          ("A — قابل للتحقيق (الموارد)","Achievable",""),
          ("R — ذو صلة (الارتباط بأهداف المركز)","Relevant",""),
          ("T — محدد زمنياً (الموعد/الإطار)","Time-bound","")]
    for i,(ar,en,dv) in enumerate(rows):
        lbl=t.rows[i].cells[1]; val=t.rows[i].cells[0]
        _set_cell_bg(lbl,LIGHT if i>0 else SWIRL); _cell_margins(lbl); _cell_margins(val,120,120,140,140)
        _valign(lbl); _valign(val)
        p=lbl.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
        run(p,ar,9.5,EVENING,True);
        vp=val.paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(vp)
        if dv: run(vp,dv,10,INK)
        else: run(vp," ",10)
    doc.add_paragraph().paragraph_format.space_after=Pt(6)
goal_card()

# Worked example
para("مثال مُعبّأ (نموذجي):",10.5,EVENING,True,space_before=2,space_after=3)
ex=doc.add_table(rows=6,cols=2); ex.alignment=WD_TABLE_ALIGNMENT.RIGHT
set_table_width(ex,[15.8,3.0]); _table_rtl(ex); _set_borders(ex,"D9C98C")
ex_rows=[("الهدف","Goal","استقطاب زوّار نوعيين لجناح المركز وتحقيق تغطية إعلامية خارجية خلال المعرض."),
    ("S — محدد","Specific","زيادة تفاعل الجمهور المتخصص في سلامة النقل مع جناح المركز."),
    ("M — قابل للقياس","Measurable","عدد زوّار يعادل 10٪ من إجمالي زوّار المعرض + تغطية إعلامية خارجية واحدة على الأقل."),
    ("A — قابل للتحقيق","Achievable","فريق من 4 ممثّلين + مواد تعريفية جاهزة + جدول لقاءات مسبق."),
    ("R — ذو صلة","Relevant","يدعم الصورة الذهنية للمركز وأهداف التسويق الإستراتيجي."),
    ("T — محدد زمنياً","Time-bound","خلال أيام المعرض (٣ أيام) مع رفع التقرير خلال ٥ أيام عمل من انتهائه.")]
for i,(ar,en,dv) in enumerate(ex_rows):
    lbl=ex.rows[i].cells[1]; val=ex.rows[i].cells[0]
    _set_cell_bg(lbl,"FBF3D9"); _cell_margins(lbl); _cell_margins(val)
    _valign(lbl); _valign(val)
    p=lbl.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
    run(p,ar,9,"8A6D12",True)
    vp=val.paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(vp)
    run(vp,dv,9.5,INK)
doc.add_paragraph().paragraph_format.space_after=Pt(6)

# ================= 5. KPI TARGETS REFERENCE =================
section_head("ربط الأهداف بمؤشرات الأداء المعتمدة","Linked KPI Targets","5")
t=doc.add_table(rows=5,cols=4); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
set_table_width(t,[7.0,4.0,3.0,4.8]); _table_rtl(t); _set_borders(t)
hd=["المؤشر · KPI","النوع · Type","المستهدف · Target","الناتج الفعلي · Actual"]
for j,h in enumerate(hd):
    cc=t.rows[0].cells[j]; _set_cell_bg(cc,EVENING); _cell_margins(cc,70,70,100,100); _valign(cc)
    pp=cc.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl_par(pp)
    run(pp,h,9.5,WHITE,True)
kpis=[("عدد زوّار الجناح","نسبي","5–10٪ من الزوّار"),
      ("مدّة بقاء الزائر","كمّي","2–5 دقائق"),
      ("الشراكات المستهدَفة","كمّي","2"),
      ("التغطيات الإعلامية الخارجية","كمّي","1")]
for i,(name,typ,tgt) in enumerate(kpis,1):
    row=t.rows[i]
    for j in range(4): _cell_margins(row.cells[j]); _valign(row.cells[j])
    if i%2==0:
        for j in range(4): _set_cell_bg(row.cells[j],"F4F6F2")
    p=row.cells[0].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p); run(p,name,10,EVENING,True)
    p=row.cells[1].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl_par(p); run(p,typ,10,HIGHLAND)
    p=row.cells[2].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _rtl_par(p); run(p,tgt,10,INK,True)
    p=row.cells[3].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(p," ",10)
doc.add_paragraph().paragraph_format.space_after=Pt(8)

# ================= 6. SIGN-OFF =================
section_head("الإعداد والاعتماد","Preparation & Approval","6")
t=doc.add_table(rows=2,cols=2); t.alignment=WD_TABLE_ALIGNMENT.RIGHT
set_table_width(t,[9.4,9.4]); _table_rtl(t); _set_borders(t)
labels=[("إعداد: مدير المشاركة","Prepared by"),("اعتماد: مدير عام الإدارة","Approved by")]
for j,(ar,en) in enumerate(labels):
    c=t.rows[0].cells[j]; _set_cell_bg(c,LIGHT); _cell_margins(c); _valign(c)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
    run(p,ar,10,EVENING,True); run(p,f"  ·  {en}",8,HIGHLAND)
for j in range(2):
    c=t.rows[1].cells[j]; _cell_margins(c,260,260,140,140); _valign(c)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; _rtl_par(p)
    run(p,"الاسم / التوقيع / التاريخ:",9,"9AA7A2",italic=True)

para("",space_after=2)
para("نقيس الأثر · نوثّق الدروس · نرتقي بالحضور",9,HIGHLAND,False,WD_ALIGN_PARAGRAPH.CENTER,space_before=6)

doc.save("NTSC-SMART-Goal-Template.docx")
print("saved NTSC-SMART-Goal-Template.docx")
